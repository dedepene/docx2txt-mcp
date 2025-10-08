# D:\Code\projects\mcp\docx2txt\server.py

from fastmcp import FastMCP # Only import FastMCP
from docx import Document
from typing import Union
import io

# 1. Initialize the FastMCP server, which automatically creates the mcp.tool decorator.
# Use 'mcp' as the variable name for compatibility with the decorator syntax.
mcp = FastMCP("DocxConverter")

# 2. Use the tool decorator from the created mcp object.
@mcp.tool
def docx_to_txt(
    file_path: Union[str, None] = None, 
    file_content_base64: Union[bytes, None] = None
) -> str:
    """
    Converts a Microsoft Word (.docx) document to plain text.
    Input can be a local file path or the base64-encoded content of the file.
    
    :param file_path: The local path to the .docx file.
    :param file_content_base64: The base64-encoded content of the .docx file.
    :return: The plain text content of the document.
    """
    if file_path:
        # Use file path if provided
        doc = Document(file_path)
    elif file_content_base64:
        # Use content if provided. FastMCP handles base64-to-bytes conversion for 'bytes' type hint.
        # Load document from a BytesIO strea                        m
        doc = Document(io.BytesIO(file_content_base64))
    else:
        raise ValueError("Either file_path or file_content_base64 must be provided.")

    text_content = []
    # Iterate through all paragraphs and append text
    for paragraph in doc.paragraphs:
        text_content.append(paragraph.text)

    # Join the text with newlines
    return '\n'.join(text_content)

if __name__ == "__main__":
    # Run the server with the default stdio transport for local integration
    mcp.run()